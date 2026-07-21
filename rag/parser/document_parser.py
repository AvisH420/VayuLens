"""Document ingestion pipeline.

Extracts text, tables, headings, and provenance metadata from PDF, DOCX, TXT,
HTML, Markdown, and scanned PDFs (OCR). Optional dependencies degrade
gracefully: if a parser lib is missing, the loader raises a clear error only
when that format is actually requested.
"""
from __future__ import annotations

import hashlib
import re
from pathlib import Path

from rag.logging_utils import get_logger
from rag.types import DocumentMetadata, ParsedDocument, ParsedElement

log = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm"}


def _doc_id(path: Path) -> str:
    return hashlib.sha1(str(path.resolve()).encode()).hexdigest()[:16]


def _infer_doc_type(name: str) -> str | None:
    n = name.lower()
    table = {
        "grap": "GRAP",
        "ncap": "NCAP",
        "cpcb": "CPCB",
        "caqm": "CAQM",
        "air act": "Air Act",
        "air_act": "Air Act",
        "environment": "Environment Protection Act",
        "factory": "Factory Act",
        "factories": "Factory Act",
        "intervention": "Intervention Report",
    }
    for key, val in table.items():
        if key in n:
            return val
    return None


def _base_metadata(path: Path) -> DocumentMetadata:
    year = None
    m = re.search(r"(19|20)\d{2}", path.stem)
    if m:
        year = int(m.group())
    return DocumentMetadata(
        source=path.name,
        doc_id=_doc_id(path),
        title=path.stem.replace("_", " ").strip(),
        doc_type=_infer_doc_type(path.name),
        year=year,
    )


def _split_paragraphs(text: str) -> list[str]:
    parts = re.split(r"\n\s*\n", text)
    return [p.strip() for p in parts if p.strip()]


def _is_heading(line: str) -> tuple[bool, int]:
    s = line.strip()
    if not s or len(s) > 120:
        return False, 0
    if s.startswith("#"):
        return True, len(s) - len(s.lstrip("#"))
    if re.match(r"^(chapter|section|part|rule|clause|schedule)\b", s, re.I):
        return True, 1
    if re.match(r"^\d+(\.\d+)*\s+\S", s) and len(s) < 90:
        return True, 2
    if s == s.upper() and len(s.split()) <= 12 and any(c.isalpha() for c in s):
        return True, 1
    return False, 0


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------
def _parse_txt(path: Path, meta: DocumentMetadata) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return _elements_from_text(text, meta, page=1)


def _parse_markdown(path: Path, meta: DocumentMetadata) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    elements: list[ParsedElement] = []
    current_section: str | None = None
    for block in _split_paragraphs(text):
        first = block.splitlines()[0]
        is_h, lvl = _is_heading(first)
        if is_h:
            title = first.lstrip("#").strip()
            current_section = title
            elements.append(
                ParsedElement(
                    text=title, element_type="heading", heading_level=lvl or 1,
                    section=title, page=1,
                )
            )
            rest = "\n".join(block.splitlines()[1:]).strip()
            if rest:
                elements.append(
                    ParsedElement(text=rest, section=current_section, page=1)
                )
        else:
            elements.append(
                ParsedElement(text=block, section=current_section, page=1)
            )
    return ParsedDocument(metadata=meta, elements=elements)


def _elements_from_text(
    text: str, meta: DocumentMetadata, page: int | None
) -> ParsedDocument:
    elements: list[ParsedElement] = []
    current_section: str | None = None
    for para in _split_paragraphs(text):
        first = para.splitlines()[0]
        is_h, lvl = _is_heading(first)
        if is_h and len(para.splitlines()) == 1:
            current_section = para.strip()
            elements.append(
                ParsedElement(
                    text=para.strip(), element_type="heading",
                    heading_level=lvl or 1, section=current_section, page=page,
                )
            )
        else:
            elements.append(
                ParsedElement(text=para, section=current_section, page=page)
            )
    return ParsedDocument(metadata=meta, elements=elements)


def _parse_html(path: Path, meta: DocumentMetadata) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        text = re.sub(r"<[^>]+>", "\n", raw)
        return _elements_from_text(text, meta, page=1)

    soup = BeautifulSoup(raw, "html.parser")
    if soup.title and soup.title.string:
        meta.title = soup.title.string.strip()
    elements: list[ParsedElement] = []
    section: str | None = None
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
        if tag.name.startswith("h"):
            section = tag.get_text(strip=True)
            elements.append(
                ParsedElement(
                    text=section, element_type="heading",
                    heading_level=int(tag.name[1]), section=section, page=1,
                )
            )
        elif tag.name == "table":
            elements.append(
                ParsedElement(
                    text=_html_table_to_text(tag), element_type="table",
                    section=section, page=1,
                )
            )
        else:
            txt = tag.get_text(" ", strip=True)
            if txt:
                elements.append(ParsedElement(text=txt, section=section, page=1))
    return ParsedDocument(metadata=meta, elements=elements)


def _html_table_to_text(tag) -> str:
    rows = []
    for tr in tag.find_all("tr"):
        cells = [c.get_text(" ", strip=True) for c in tr.find_all(["td", "th"])]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _parse_docx(path: Path, meta: DocumentMetadata) -> ParsedDocument:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "python-docx not installed. Install with: pip install python-docx"
        ) from e
    document = docx.Document(str(path))
    elements: list[ParsedElement] = []
    section: str | None = None
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style:
            lvl = int("".join(filter(str.isdigit, style)) or 1)
            section = text
            elements.append(
                ParsedElement(
                    text=text, element_type="heading", heading_level=lvl,
                    section=section, page=1,
                )
            )
        else:
            elements.append(ParsedElement(text=text, section=section, page=1))
    for table in document.tables:
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
        ]
        elements.append(
            ParsedElement(
                text="\n".join(rows), element_type="table", section=section, page=1
            )
        )
    return ParsedDocument(metadata=meta, elements=elements)


def _parse_pdf(path: Path, meta: DocumentMetadata, ocr_enabled: bool,
               ocr_lang: str) -> ParsedDocument:
    elements: list[ParsedElement] = []
    text_found = False
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                text_found = True
                sub = _elements_from_text(page_text, meta, page=page_num)
                elements.extend(sub.elements)
    except ImportError:
        log.warning("pypdf not installed; attempting OCR-only path for %s", path.name)

    if not text_found and ocr_enabled:
        elements.extend(_ocr_pdf(path, ocr_lang))

    if not elements:
        raise RuntimeError(
            f"Could not extract text from {path.name}. Install pypdf and/or OCR "
            "extras (pytesseract, pdf2image)."
        )
    return ParsedDocument(metadata=meta, elements=elements)


def _ocr_pdf(path: Path, ocr_lang: str) -> list[ParsedElement]:
    try:
        import pytesseract  # type: ignore
        from pdf2image import convert_from_path  # type: ignore
    except ImportError:
        log.warning("OCR extras not installed; skipping OCR for %s", path.name)
        return []
    elements: list[ParsedElement] = []
    images = convert_from_path(str(path))
    for page_num, img in enumerate(images, start=1):
        text = pytesseract.image_to_string(img, lang=ocr_lang).strip()
        if text:
            for para in _split_paragraphs(text):
                elements.append(ParsedElement(text=para, page=page_num))
    return elements


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def parse_document(
    path: str | Path,
    ocr_enabled: bool = True,
    ocr_language: str = "eng",
) -> ParsedDocument:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    meta = _base_metadata(path)
    ext = path.suffix.lower()
    log.info("Parsing %s (type=%s)", path.name, meta.doc_type)
    if ext == ".txt":
        return _parse_txt(path, meta)
    if ext in {".md", ".markdown"}:
        return _parse_markdown(path, meta)
    if ext in {".html", ".htm"}:
        return _parse_html(path, meta)
    if ext == ".docx":
        return _parse_docx(path, meta)
    if ext == ".pdf":
        return _parse_pdf(path, meta, ocr_enabled, ocr_language)
    raise ValueError(f"Unsupported file type: {ext}")


def parse_directory(
    directory: str | Path,
    ocr_enabled: bool = True,
    ocr_language: str = "eng",
) -> list[ParsedDocument]:
    directory = Path(directory)
    docs: list[ParsedDocument] = []
    for path in sorted(directory.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                docs.append(parse_document(path, ocr_enabled, ocr_language))
            except Exception as e:  # noqa: BLE001 - keep ingesting the rest
                log.error("Failed to parse %s: %s", path.name, e)
    return docs
